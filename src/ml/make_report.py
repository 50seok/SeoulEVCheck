from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
ROOT=Path("C:/teamwork/SeoulEVCheck"); FIG=ROOT/"reports"/"figures"
doc=Document()
s=doc.styles['Normal']; s.font.name='맑은 고딕'; s.font.size=Pt(10)
rf=s.element.get_or_add_rPr().get_or_add_rFonts()
rf.set(qn('w:eastAsia'),'맑은 고딕'); rf.set(qn('w:ascii'),'맑은 고딕')
def fig(n,w=5.2):
    p=FIG/n
    if p.exists(): doc.add_picture(str(p),width=Inches(w))
def bullets(items):
    for t in items: doc.add_paragraph(t, style='List Bullet')
doc.add_heading('머신러닝 산출물 — SeoulEVCheck',0)
doc.add_paragraph('서울 전기차 충전 수요 예측 및 시각화  |  담당: 오영석  |  1주차 ML')
doc.add_heading('1. 사업과제',1)
doc.add_paragraph('AI 기반 서울시 전기차 충전 수요 예측 및 시각화 (SeoulEVCheck)')
doc.add_heading('2. 개요 및 현황',1)
doc.add_heading('2.1 추진배경 및 목적',2)
bullets(['전기차 보급 급증 → 충전 인프라 수요 예측·관리 필요성 증대',
         '활용: ①충전 인프라 투자 효율 ②전력 부하 관리(에너지 절감) ③운영 최적화',
         '향후 비전: 예측 수요 기반 충전 인프라 투자 우선순위 결정 → 충전 어드바이저(LLM·RAG)로 발전'])
doc.add_heading('2.2 과제 범위',2)
# (대분류, 중분류, 소분류, 내용) — 소분류 '' 이면 중분류 옆 내용 칸을 가로 병합
scope_rows = [
    ('데이터셋',    '전처리',        '데이터 수집',       'KEPCO 서울 충전 데이터 527,183건 수집 (2025-01~2026-03)'),
    ('데이터셋',    '전처리',        '데이터 정제',       '음수·누락 등 불필요 데이터 제거 → 112,453건 정제'),
    ('데이터셋',    '전처리',        '시군구 처리',       '시군구 컬럼 직접 활용 → 25개 자치구 분류'),
    ('데이터셋',    'EDA 분석',      '분포 분석',         '충전량 분포 확인 (hist)'),
    ('데이터셋',    'EDA 분석',      '상관관계 분석',     '항목 간 상관계수 산출 및 누수 특성 식별 (heatmap)'),
    ('데이터셋',    'EDA 분석',      '구별 분석',         '수요 상위 구 파악 및 시각화 (barplot)'),
    ('데이터셋',    'EDA 분석',      '충전기 유형 분석',  '급속·완속 평균 충전량 비교'),
    ('AI 머신러닝', '모델 구축',     '',                  'XGBoost 기반 일별 충전량 예측 모델 설계'),
    ('AI 머신러닝', '학습 및 예측',  '',                  '충전 특성 선택, 8:2 데이터 분리, 예측 및 결과 출력'),
    ('AI 머신러닝', '분석',          '오차 분석',         'Actual vs Predicted 차이 분석 (scatter)'),
    ('AI 머신러닝', '분석',          '중요도 분석',       '충전량 예측 기여도 인자 분석 (barplot)'),
    ('AI 머신러닝', '모델 성능평가', '지표 계산',         'R², MAE 기반 모델 신뢰도 검증 (5회 교차검증)'),
    ('AI 머신러닝', '프로토타이핑',  '',                  'Streamlit을 이용한 충전 수요 예측 웹 앱 개발'),
]
st = doc.add_table(rows=len(scope_rows)+1, cols=4)
st.style = 'Table Grid'
# 헤더: 텍스트 먼저 설정 후 병합
hdr = st.rows[0]
hdr.cells[0].text = '과제구분'
hdr.cells[3].text = '내용'
hdr.cells[0].merge(hdr.cells[2])
# 데이터 행
for i, (cat, mid, sub, cont) in enumerate(scope_rows):
    ri = i + 1
    row = st.rows[ri]
    row.cells[0].text = cat
    row.cells[1].text = mid
    if sub:
        row.cells[2].text = sub
        row.cells[3].text = cont
    else:
        row.cells[2].text = cont
        row.cells[2].merge(row.cells[3])
# 세로 병합 (마지막에)
st.cell(1,0).merge(st.cell(7,0))    # 데이터셋
st.cell(8,0).merge(st.cell(13,0))   # AI 머신러닝
st.cell(1,1).merge(st.cell(3,1))    # 전처리
st.cell(4,1).merge(st.cell(7,1))    # EDA 분석
st.cell(10,1).merge(st.cell(11,1))  # 분석
doc.add_heading('2.3 과제 추진 방법',2)

def bpara(text):
    p=doc.add_paragraph(); p.add_run(text).bold=True; return p
def cbullet(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Pt(12); p.add_run('○ '+text); return p
def sbullet(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Pt(24); p.add_run('■ '+text); return p

bpara('1) 구축 대상 선정 기준')
cbullet('데이터 접근성 및 활용성')
sbullet('공공데이터포털을 통한 한국전력공사 서울시 전기차 충전 데이터 무료 수집 가능')
sbullet('정부 및 공공기관에서 이미 구축된 충전 이용 데이터베이스 활용 가능')
sbullet('충전량에 영향을 미치는 자치구·충전기 유형·시간 등 다양한 독립변수 포함으로 모델학습 유용성 확보')
cbullet('예측모델 개발 효율성')
sbullet('충전량·충전 구분·자치구·요일·월 등 변수 구조가 단순하여 모델 학습 및 평가 과정 간소화 가능')
sbullet('개발된 모델을 타 지역 및 전국 단위 충전 인프라 예측에 확장 적용 가능')
cbullet('환경문제 해결 기여도 및 경제성')
sbullet('충전 수요 예측을 통한 전력 부하 분산 및 에너지 절감 기여 (ex. 전력 피크 저감, 재생에너지 연계 등)')
sbullet('충전 인프라 투자 우선순위 결정을 통한 설치 비용 절감 효과')
sbullet('충전 병목 해소로 전기차 보급 확대 → 온실가스 감소 등 사회적 비용감소 효과')

bpara('2) AI 예측 분석모델 적용 대상')
at=doc.add_table(rows=2,cols=4); at.style='Table Grid'
for j,h in enumerate(['환경관리\n기능','수집 데이터','예측모델인자(독립변수)','AI예측 분석 대상']):
    at.rows[0].cells[j].text=h
dr=at.rows[1]
dr.cells[0].text='전기차\n충전'
dr.cells[1].text=('- 한국전력공사 서울시 전기차 충전소 충전량 데이터 (527,183건, 2025-01~2026-03)\n'
                   '- 충전소별 일별 충전량 집계 데이터\n'
                   '- 자치구별 일별 충전량 집계 데이터')
dr.cells[2].text=('- 위치변수: 자치구, 충전소명\n'
                   '- 충전기변수: 충전 구분(급속/완속), 충전기 용량(kW)\n'
                   '- 시간변수: 요일, 월')
dr.cells[3].text=('- 일별 충전량(kWh) 예측\n'
                   '- 수요 상위 자치구 도출\n'
                   '- 충전소 핫스팟 TOP10\n'
                   '- 인프라 투자 우선순위 결정')
doc.add_heading('3. 연구개발 주요 결과물',1)
doc.add_heading('① 데이터 수집',2)
doc.add_paragraph('한국전력공사 서울 충전량, 527,183 세션, 19개 컬럼(충전량·충전구분·시군구·날짜 등). 기간: 2025-01~2026-03. 출처: KEPCO 공공데이터 활용승인.')
doc.add_heading('② 데이터 분석',2)
doc.add_paragraph('데이터 정제: 음수·누락 등 불필요 데이터 제거 · 분석 기간(2025-01~2026-03) · 시군구 컬럼으로 자치구 직접 분류 (25개 구 100%) · 충전 후에야 알 수 있는 항목(충전 횟수) 예측에서 제외.')
def comment(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x33,0x33,0x55)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'EEF4FB')
    p._p.get_or_add_pPr().append(shd)

doc.add_paragraph('항목 간 관련도 분석 — 충전 후에야 알 수 있는 항목은 관련도가 높아도 학습에서 제외:'); fig('06_corr.png',4.3)
comment('세션수(0.89)·avg_hour(0.48)는 타깃과 높은 상관을 보이나, 충전 완료 후에야 알 수 있는 누수 특성이므로 학습에서 제외. 연도·월·요일은 사전 알 수 있는 독립변수로 선택.')

doc.add_paragraph('충전량 분포(0~4000 확대) / 충전구분별 평균:'); fig('01_target_dist.png',4.5); fig('05_charger.png',3.3)
comment('충전량 분포는 0~500 kWh 구간에 집중된 우편향 분포 → log1p 변환으로 정규화 후 학습. 급속(685 kWh)이 완속(525 kWh)보다 약 1.3배 높아 충전기 유형이 주요 변수임을 확인.')

doc.add_paragraph('수요 상위 자치구역(인프라 투자 우선 후보 지역):'); fig('04_top_gu.png',5.0)
comment('송파·강남·서초 순으로 수요 집중. 이상치(서울숲M타워) 제거 후 성동구는 3위→5위로 조정됨. 수요 상위 구를 인프라 확충 우선순위로 활용.')

doc.add_paragraph('요일별 충전량 추이 — 요일(weekday)을 독립변수로 선택한 근거:'); fig('02_weekday.png',5.0)
comment('평일(월~금) 대비 주말(토·일) 충전량이 낮은 경향 확인 → 업무지구 주간 충전 패턴. 요일별 차이가 유의미하여 weekday를 독립변수로 채택.')

doc.add_paragraph('월별 충전량 추이 — 월(month)을 독립변수로 선택한 근거:'); fig('03_month.png',5.0)
comment('하반기(7~10월)에 충전량 피크 → 여름철 냉방 연계 및 전기차 이용 증가 영향. 계절성이 뚜렷해 month를 독립변수로 채택하고 month_seq로 장기 트렌드도 반영.')
doc.add_heading('③ 데이터 학습 및 모델 정의',2)

def code_para(text):
    p=doc.add_paragraph()
    p.paragraph_format.left_indent=Pt(12)
    r=p.add_run(text)
    r.font.name='Courier New'; r.font.size=Pt(9)
    r.font.color.rgb=RGBColor(0x33,0x33,0x33)
    shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F2F2F2')
    p._p.get_or_add_pPr().append(shd)

def out_para(text):
    p=doc.add_paragraph()
    p.paragraph_format.left_indent=Pt(12)
    r=p.add_run(text)
    r.font.name='Courier New'; r.font.size=Pt(9)
    r.font.color.rgb=RGBColor(0x44,0x44,0x44)
    shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'FAFAFA')
    p._p.get_or_add_pPr().append(shd)

# 3.1 모델 정의
p=doc.add_paragraph(); p.add_run('3.1  모델 정의 및 컴파일').bold=True
p=doc.add_paragraph(); p.add_run('○ 모델 정의 : RandomForest').bold=True
code_para('# RandomForest 회귀 모델 정의')
code_para('from sklearn.ensemble import RandomForestRegressor')
code_para('model = RandomForestRegressor(n_estimators=50, random_state=42, n_jobs=-1)')
p=doc.add_paragraph(); p.add_run('○ 모델 컴파일 / 학습').bold=True
code_para('X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)')
code_para('model.fit(X_train, y_train)')

# 3.2 모델 학습 및 시각화
p=doc.add_paragraph(); p.add_run('3.2  모델학습 및 학습 시각화').bold=True
p=doc.add_paragraph(); p.add_run('○ 모델 학습 결과').bold=True
out_para('→ 최적 모델 : RandomForest   R²=0.940   RMSE=2,533   MAE=1,697')
out_para('   CV R²(3-fold) : 0.951 ± 0.007  |  학습 데이터 : 750행 (월별 집계)')
out_para('   이상치 9건 제거 후 R² 0.862 → 0.940 향상')

def styled_table(doc, headers, rows, hdr_fill='2E5F8A', hdr_color='FFFFFF',
                 alt_fill='EBF3FB', border_color='2E5F8A'):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    def set_cell_bg(cell, fill):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'), fill)
        cell._tc.get_or_add_tcPr().append(shd)
    def set_cell_border(cell, color):
        tcp = cell._tc.get_or_add_tcPr()
        for side in ['top','left','bottom','right']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
            b.set(qn('w:space'),'0'); b.set(qn('w:color'), color)
            bdr = OxmlElement('w:tcBorders')
            bdr.append(b)
            tcp.append(bdr)
    # 헤더 행
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, hdr_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True; run.font.color.rgb = RGBColor.from_string(hdr_color)
        run.font.size = Pt(9.5)
    # 데이터 행
    for i, row in enumerate(rows):
        fill = alt_fill if i % 2 == 0 else 'FFFFFF'
        for j, v in enumerate(row):
            cell = t.rows[i+1].cells[j]
            cell.text = v
            set_cell_bg(cell, fill)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].font.size = Pt(9.5)
    return t

# 5개 모델 비교 표
p=doc.add_paragraph(); p.add_run('○ 모델 비교 (구 단위, 월별 집계 기준)').bold=True
styled_table(doc,
    ['모델', 'R²', 'RMSE(kWh)', 'MAE(kWh)', 'CV R²(3-fold)', '비고'],
    [['LinearRegression',  '0.676', '5,876', '4,652', '0.682±0.036', ''],
     ['Ridge',             '0.691', '5,734', '4,570', '0.680±0.033', ''],
     ['DecisionTree',      '0.463', '7,561', '5,705', '0.617±0.119', ''],
     ['RandomForest',      '0.940', '2,533', '1,697', '0.951±0.007', '★ 최적'],
     ['GradientBoosting',  '0.740', '5,266', '3,999', '0.822±0.058', '']],
    hdr_fill='2E5F8A')
doc.add_paragraph('→ RandomForest가 R²=0.940으로 전 모델 중 최고 성능. 소규모 월별 집계 데이터에서 과적합 없이 안정적.')

# 다중공선성 VIF 분석
p=doc.add_paragraph(); p.add_run('○ 다중공선성(VIF) 분석').bold=True
doc.add_paragraph('수치형 피처(month_seq·year·month·avg_hour) 간 다중공선성을 VIF(분산팽창인수)로 검증.')
styled_table(doc,
    ['변수', 'VIF', '판정', '조치'],
    [['month_seq', '4.8',   '양호 (<5)',   'year+month 선형결합 → 제거 또는 단독 사용'],
     ['year',      '173.2', '심각 (>10)',  'month_seq와 중복 → LinearRegression 성능 저하 원인'],
     ['month',     '3.7',   '양호 (<5)',   'month_seq와 독립적 계절성 반영 → 유지'],
     ['avg_hour',  '1.0',   '양호 (<5)',   '독립적 → 유지']],
    hdr_fill='2E5F8A')
doc.add_paragraph('→ year(VIF=173.2) 심각 — month_seq에 이미 포함된 정보. RandomForest는 다중공선성 면역이므로 성능 무영향. LinearRegression R²=0.676 저조의 주요 원인.')
fig('13_vif_analysis.png', 6.0)

p=doc.add_paragraph(); p.add_run('○ 학습 과정 시각화 — 트리 수 증가에 따른 R² 변화').bold=True
doc.add_paragraph('트리 수(n_estimators)를 1→200까지 늘리며 학습·테스트 R² 변화를 추적. n=50 이후 성능 수렴, 학습/테스트 격차 좁아 과적합 없음 확인.')
fig('14_learning_curve.png', 6.0)

# 최종 선정 모델 성능 요약
p=doc.add_paragraph(); p.add_run('○ 최종 선정 모델 성능 요약').bold=True
styled_table(doc,
    ['모델', '기준 모델 R²', 'AI 모델 R²', 'CV R²(3-fold)', 'RMSE(kWh)'],
    [['구(거시) — RandomForest',  '0.277', '0.940', '0.951±0.007', '2,533'],
     ['충전소(미시) — LinearRegression', '0.083', '0.058', '0.369±0.050', '277']])
doc.add_paragraph('→ 구 모델은 기준 모델 대비 성능 대폭 향상 확인.')

p=doc.add_paragraph(); p.add_run('3.3  모델 예측 및 평가').bold=True
p=doc.add_paragraph(); p.add_run('○ 예측값 vs 실제값 / 특성 중요도').bold=True
it=doc.add_table(rows=2,cols=2); it.style='Table Grid'
for ci,(fn,cap,cmt) in enumerate([
    ('07_pred_gu.png',
     '[그림] 예측값 vs 실제값 산점도',
     '대부분의 점이 대각선(y=x)에 밀집 → 예측값이 실제값과 높은 일치도를 보임.\nR²=0.94로 실제 충전량 변동의 94%를 모델이 설명함.'),
    ('08_imp_gu.png',
     '[그림] 특성 중요도 TOP12',
     '금천구(gu_금천구)가 1위 — 월평균 10,008 kWh로 25개 구 최하위의 극단적 패턴 반영.\nmonth_seq(시계열)·month(계절성)가 상위권 → 충전 수요는 시간 흐름에 따라 지속 증가.')]):
    cell = it.rows[0].cells[ci]
    cp = cell.add_paragraph()
    if (FIG/fn).exists():
        cp.add_run().add_picture(str(FIG/fn), width=Inches(3.0))
    cell.add_paragraph(cap)
    c_cell = it.rows[1].cells[ci]
    cp2 = c_cell.add_paragraph()
    r = cp2.add_run(cmt)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x33,0x33,0x55)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'EEF4FB')
    c_cell._tc.get_or_add_tcPr().append(shd)
doc.add_heading('④ 예측 결과 활용 방안',2)
doc.add_paragraph(
    '충전 수요 예측 결과는 단순 수치가 아니라 \'어디에, 언제, 얼마나\' 충전 자원을 투입할지 '
    '의사결정의 근거로 활용된다. 아래 표는 예측 아웃풋의 3가지 실사용 경로를 정리한 것이다.'
)
ut=doc.add_table(rows=4,cols=3); ut.style='Table Grid'
uhdr=['활용 분야','예측 활용 방식','기대 효과']
udata=[
    ['충전 인프라 투자','수요 상위 구·충전소 TOP-N 파악','신규 충전기 설치 우선순위 결정 (예산 효율 향상)'],
    ['전력 부하 관리','시간대·요일별 수요 피크 예측','한전 부하 분산 계획 수립 (에너지 절감)'],
    ['이동형 충전 배차(비전)','고수요 구역·시간대 사전 탐지','이동 충전 차량 출동 지역·시각 결정'],
]
for j,h in enumerate(uhdr): ut.rows[0].cells[j].text=h
for i,r in enumerate(udata):
    for j,v in enumerate(r): ut.rows[i+1].cells[j].text=v
doc.add_paragraph('')
doc.add_paragraph(
    '【활용 예시】 앱에서 강남구 / 급속 / 금요일 / 6월 선택 시 예측 충전량 1,245 kWh 출력 '
    '→ 이 구역·시간대는 수요 상위권 → 인프라 투자 및 이동형 충전 우선 배치 대상으로 판단.'
)
doc.add_heading('⑤ 예측 모델 실효성 검증 — 25년 학습 모델 vs 26년 1~3월 실제 충전량',2)
doc.add_paragraph(
    '25년(1~12월) 데이터로 학습한 모델이 26년 1~3월 실제 충전량을 얼마나 정확하게 예측하는지 검증하였다. '
    '학습(25년)·검증(26년 Q1) 분리를 통해 모델의 미래 예측력을 확인한다.'
)
vt=doc.add_table(rows=7,cols=4); vt.style='Table Grid'
vhdr=['구분','내용','수치','비고']
vdata=[
    ['학습 데이터','25년 1~12월 충전 세션','서울 527,183건','KEPCO 전체'],
    ['검증 데이터','26년 1~3월 충전 세션','서울 106,569건','out-of-sample'],
    ['구 모델 R²','실제 vs 예측 일치도','0.806','RandomForest'],
    ['충전소 모델 R²','실제 vs 예측 일치도','0.564','RandomForest'],
    ['수요 1위 구','25년 충전량 최다','송파구 252,564 kWh','26년 투자 우선'],
    ['수요 2위 구','25년 충전량 2위','강남구 244,968 kWh','인프라 확충 필요'],
]
for j,h in enumerate(vhdr): vt.rows[0].cells[j].text=h
for i,r in enumerate(vdata):
    for j,v in enumerate(r): vt.rows[i+1].cells[j].text=v
doc.add_paragraph(
    '→ 25년 데이터 학습 모델은 26년 1~3월 실제 충전량을 R²=0.806(구 모델) 수준으로 예측하여 '
    '인프라 투자 의사결정 근거로 활용 가능함을 검증하였다. '
    '수요 상위 구(송파·강남·성동·서초)를 중심으로 26년 충전 인프라 확충이 필요하다.'
)
doc.add_heading('⑥ 프로토타이핑 (화면)',2)
doc.add_paragraph('Streamlit 서비스: 구·충전기·요일·월 입력 → 예상 일 충전량 + 구별 수요 + 충전소 핫스팟 TOP10.')
doc.add_paragraph('[여기에 배포된 Streamlit 앱 스크린샷을 삽입하세요]')
doc.add_heading('4. 인프라 투자 우선순위 분석',1)
doc.add_paragraph(
    '3장 모델 예측 결과를 바탕으로 구별 충전 수요와 현재 충전소 공급량을 비교하여 '
    '인프라 투자 우선순위를 도출하였다. '
    '충전소 1개당 일평균 부담 수요(kWh)를 핵심 지표로 사용하며, '
    '수요·공급 중앙값 기준 4분면 매트릭스로 투자 유형을 분류한다.'
)
doc.add_heading('4.1 갭 분석 방법론',2)
doc.add_paragraph('분석 지표: 충전소 1개당 일평균 부담 수요 = 구별 일평균 충전량(kWh) ÷ 충전소 수')
styled_table(doc,
    ['사분면', '조건', '의미', '권고'],
    [['긴급 증설', '수요↑ + 공급↓', '충전소 1개당 부하 과중', '즉시 증설 필요'],
     ['현상 유지', '수요↑ + 공급↑', '수요·공급 균형', '현 수준 유지·모니터링'],
     ['모니터링', '수요↓ + 공급↓', '전기차 보급 확대 시 수요 증가 가능', '중장기 준비'],
     ['과잉 공급', '수요↓ + 공급↑', '충전소 이용률 저조', '타 지역 재배치 검토']],
    hdr_fill='2E5F8A')

doc.add_heading('4.2 수요-공급 갭 분석 결과',2)
fig('15_infra_gap.png', 6.2)
comment('왼쪽: 충전소 1개당 부담 수요 순위 — 광진구(16.5), 강북구(14.1), 금천구(12.7) 순으로 인프라 부족이 심각. '
        '오른쪽: 수요 vs 공급 4분면 — 광진구·강북구·금천구·서대문구는 수요↑ + 공급↓ 영역에 위치하여 긴급 증설 대상.')

doc.add_heading('4.3 긴급 증설 대상 구역 (상위 4개구)',2)
styled_table(doc,
    ['자치구', '일평균 수요(kWh)', '충전소 수(개)', '충전소당 부담(kWh)', '분류'],
    [['광진구', '132.4', '8',  '16.55', '긴급 증설'],
     ['강북구', '140.6', '10', '14.06', '긴급 증설'],
     ['금천구', '76.1',  '6',  '12.68', '긴급 증설'],
     ['서대문구', '76.6', '14', '5.47', '긴급 증설']],
    hdr_fill='C0392B')
doc.add_paragraph(
    '→ 광진구·강북구는 충전 수요 자체가 평균 대비 2배 이상이며 충전소도 10개 미만으로 병목이 가장 심각하다. '
    '금천구는 수요는 중간 수준이나 충전소 6개로 서울 최저 수준이어서 즉각적인 증설이 필요하다.'
)

doc.add_heading('5. 결론 및 향후',1)
bullets(['구 모델 정확도 80.6%, 충전소 모델 정확도 56.4% — 두 모델 모두 기준 모델 대비 성능 향상 확인',
         '수요 상위 지역(송파·강남·성동·서초…) 도출 → 26년 인프라 투자 우선순위 의사결정 지원',
         '갭 분석 결과: 광진구·강북구·금천구 긴급 증설 / 강남구·서초구·송파구 과잉 공급 → 재배치 검토',
         '25년 학습 → 26년 Q1 검증 완료 — 미래 수요 예측 실효성 입증',
         '향후: 2주 DL(LSTM 시계열 예측) / 3주 LLM(충전 어드바이저·RAG 챗봇)으로 확장'])
out=ROOT/'docs'/'오영석_머신러닝프로젝트.docx'
try:
    doc.save(str(out)); print("docx 저장:", out.name)
except PermissionError:
    alt=ROOT/'docs'/'오영석_머신러닝프로젝트_new.docx'; doc.save(str(alt))
    print("원본이 열려있어 새 파일로 저장:", alt.name, "(기존 닫고 이걸 쓰세요)")
